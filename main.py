import re
import json
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
import io
from openai import OpenAI
from colorama import init, Fore
from dotenv import load_dotenv

init(autoreset=True)
load_dotenv()

client = OpenAI(api_key=os.getenv('API_TOKEN_CHAT_GPT'))
COUNT_MONEY = 0

def read_text_file(file_path):
    """
    Reads a text file and returns its content with spaces and line breaks preserved.

    :param file_path: The path to the text file.
    :return: The content of the file as a string.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        return "File not found."
    except Exception as e:
        return str(e)

def parse_text_to_json(text):
    """
    The function `parse_text_to_json` takes a text input and parses it into a JSON format, separating
    chapters and their titles.
    
    :param text: The `text` parameter is a string that represents the text to be parsed into JSON
    format. This text should contain chapter titles in the format "Chapter {chapter_number}:
    {chapter_title}", followed by the content of each chapter. Each chapter should be separated by a new
    line
    :return: a JSON string representation of the parsed text.
    """
    lines = text.split('\n')

    # Regular expression pattern for chapter titles
    chapter_regex = r'Chapter (\d+): (.+)'
    json_data = []
    current_chapter = None

    for line in lines:
        chapter_match = re.match(chapter_regex, line)

        if chapter_match:
            # Add the previous chapter to json_data
            if current_chapter is not None:
                json_data.append(current_chapter)

            chapter_number, chapter_title = chapter_match.groups()
            current_chapter = {"chapter": f"Chapter {chapter_number}", "title": chapter_title, "data": []}
        elif current_chapter is not None:
            # Add other lines as data to the current chapter
            if line.strip():
                current_chapter["data"].append(line.strip())

    # Add the last chapter
    if current_chapter is not None:
        json_data.append(current_chapter)

    return json.dumps(json_data, indent=4)

def chatgpt4_query(prompt, model="gpt-4"):
    """
    Send a query to the OpenAI GPT-4 API.

    :param prompt: The prompt to send to the model.
    :param model: The model to use, default is GPT-4.
    :param api_key: Your OpenAI API key.
    :return: The response from the model.
    """
    global COUNT_MONEY
    def calculate_cost(usage):
        input_cost_per_thousand = 0.03
        output_cost_per_thousand = 0.06

        input_tokens = usage.prompt_tokens
        output_tokens = usage.completion_tokens 

        total_cost = (input_tokens / 1000 * input_cost_per_thousand) + (output_tokens / 1000 * output_cost_per_thousand)
        return total_cost

    while True:
        try:        
            response = client.chat.completions.create(
            model=model,
            messages=[
                {
                "role": "user",
                "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=5000,
            top_p=1
            )

            cost = calculate_cost(response.usage)
            COUNT_MONEY += cost
            return str(response.choices[0].message.content)
        except Exception as e:
            print(e)
            continue
    
def text_gpt(title, text, COMMENTS):
    template = f"""{text}\n\nWrite this section of the ebook "{title}" about 300 words, don't use bullets or subsections, write it smoothly as one section. {COMMENTS}"""
    result = chatgpt4_query(template)
    result = str(result).replace(title, "")
    return result

def generate_e_book(text_array, text, COMMENTS):
    """
    The `generate_e_book` function takes in an array of text data and generates an e-book in Word format
    with chapters, titles, and body text.
    
    :param text_array: The `text_array` parameter is a list of dictionaries. Each dictionary represents
    a chapter in the e-book and contains the following keys:
    :param text: The "text" parameter is the input text that will be used as a prompt for generating the
    content of the e-book. It will be passed to the "text_gpt" function, which is not provided in the
    code snippet. The "text_gpt" function is expected to generate a segment
    :return: a byte buffer containing the generated e-book.
    """
    # Create a new Word document
    time_start_e_book = time.time()
    document = Document()

    try:
        iteration_count = 0
        for text_box in text_array:
            if text_box['data']:
                iteration_count += len(text_box['data'])
        
        print(f"{Fore.GREEN}[ОБЩЯЯ ИНФОРМАЦИЯ ОБ E-BOOK][ГЛАВ: {len(text_array)}][ИТЕРАЦИЙ ВСЕГО: {iteration_count}]")
    except Exception as e:
        print(e)
    
    count_iter_f = 0
    for index, text_box in enumerate(text_array):
        # Add a page break for each new chapter except the first one
        if document.paragraphs:
            document.add_page_break()

        # Add chapter title in bold and centered
        chapter = document.add_paragraph()
        run = chapter.add_run(text_box['chapter'])
        run.bold = True
        chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the title, centered
        title = document.add_paragraph()
        title_run = title.add_run(text_box['title'])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process the body of the chapter
        body = ""
        if text_box['data']:
            for idx, item in enumerate(text_box['data']):
                count_iter_f += 1
                index_txt = index + 1
                idx_txt = idx + 1
                time_start = time.time()
                print(f"[{count_iter_f}][ГЛАВА {index_txt}/{len(text_array)}][ПОДСЕКЦИЯ {idx_txt}/{len(text_box['data'])}] Генерация текста \t| {Fore.BLUE}WAIT...{Fore.WHITE}")
                segment = text_gpt(item, text, COMMENTS)
                time_end = time.time()
                df = time_end - time_start
                print(f"[{count_iter_f}][ГЛАВА {index_txt}/{len(text_array)}][ПОДСЕКЦИЯ {idx_txt}/{len(text_box['data'])}] Генерация текста \t| {Fore.GREEN}DONE!{Fore.WHITE} | {Fore.YELLOW} sec: {df:.2f}s {Fore.WHITE}")
                if segment:
                    body += segment + "\n"

        # Add the body text
        document.add_paragraph(body)

    # Save the document to a byte buffer
    time_end_e_book = time.time()
    df_count = time_end_e_book - time_start_e_book
    print(f"Общее время создании книги {df_count:.2f} | Потрачено денег в $: {COUNT_MONEY:.2f}")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    COMMENTS = input("[BOT] If you want to add some commentary-command to the generation, write it, otherwise just press Enter leaving the line empty\n")
    # Get text
    text = read_text_file(os.path.join(os.getcwd(), "text.txt"))
    # Get json
    parsed_json = parse_text_to_json(text)
    # Generete e-book
    e_book_buffer = generate_e_book(json.loads(parsed_json), text, COMMENTS)
    # Save the buffer to a file
    with open(os.path.join(os.getcwd(), "generated_e_book.docx"), "wb") as f:
        f.write(e_book_buffer.read())
    # Open Word
    try:
        os.startfile(os.path.join(os.getcwd(), "generated_e_book.docx"))
    except:
        pass